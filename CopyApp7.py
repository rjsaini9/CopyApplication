import os
import threading
import tkinter
import tkinter.messagebox
import pyautogui
import docx
from docx.shared import Inches, Pt
import LastResort2
import LastResort3
from datetime import date
import tkinter.font as font
from pynput import keyboard


global imageCount
imageCount = 0
global textCount
textCount = 0

top = tkinter.Tk()
top.title = "CopyMe"
top.geometry('130x80')
top.resizable(True, True)
top.attributes("-topmost", True)
top.focus()


global doc
doc = docx.Document()
section = doc.sections[0]
header = section.header
t = header.add_table(4, 2, width=Inches(9.0))
t.style = 'TableGrid'
tday = date.today()
strToday = tday.strftime("%d %b %Y")
t.rows[0].cells[0].text = "Tester"
t.rows[0].cells[1].text = "Environment"
t.rows[1].cells[0].text = "Story"
t.rows[1].cells[1].text = "Timestamp = " + strToday
t.rows[2].cells[0].text = "User"
t.rows[2].cells[1].text = "Role"
t.rows[3].cells[0].text = "Location"
t.rows[3].cells[1].text = "Browser"


def method3(sf,commenttext):
    global doc
    if os.path.exists(filepath +"\screen.docx")==True:
        print ("exists")
    else:
        doc = None
        doc = docx.Document()
        section = doc.sections[0]
        header = section.header
        t = header.add_table(4, 2, width=Inches(9.0))
        t.style = 'TableGrid'
        tday = date.today()
        strToday = tday.strftime("%d %b %Y")
        t.rows[0].cells[0].text = "Tester"
        t.rows[0].cells[1].text = "Environment"
        t.rows[1].cells[0].text = "Story"
        t.rows[1].cells[1].text = "Timestamp = " + strToday
        t.rows[2].cells[0].text = "User"
        t.rows[2].cells[1].text = "Role"
        t.rows[3].cells[0].text = "Location"
        t.rows[3].cells[1].text = "Browser"

    p = doc.add_paragraph()
    r= p.add_run()
    r.add_break()
    global textCount
    textCount = textCount + 1
    textCountStr = str(textCount) + ".  "
    commenttext = textCountStr + commenttext
    r.add_text(commenttext)
    r.add_picture(filepath+"\screen.png", width= Inches(6), height= Inches(4))
    global imageCount
    imageCount = imageCount + 1
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    r.add_text("\t\t\t\t\t Figure "+ str(imageCount))
    doc.save(filepath+"\screen.docx")
    sf.withdraw()
    sf.destroy()
    top.deiconify()

def method1():
    sf = tkinter.Tk()
    sf.geometry('100x80')
    sf.resizable(True,True)
    sf.attributes("-topmost",True)
    entrySF = tkinter.Entry(sf,bd= 1,width= 12)
    entrySF.place(x=30, y =25)
    lblEnterDEsc= tkinter.Label(sf,text="Enter Desc").place(x=8,y=0)
    btnOk= tkinter.Button(sf,text="Ok",command= lambda:  method3(sf, entrySF.get()))
    btnOk.place(x=40,y=50)
    sf.deiconify
    #while True:
    #    sf.update()
    #sf.mainloop()

    #sf.focus_force()
    #entrySF.focus_set()
    print ("methd1")

def helloCallBack():
    global filepath
    top.withdraw()
    top.iconify()
    screenshot= pyautogui.screenshot()
    filepath= entryW.get()
    screenshot.save(filepath+"\screen.png")
    print (filepath)
    sf = tkinter.Tk()
    sf.geometry('100x80')
    sf.resizable(True,True)
    sf.attributes("-topmost",True)
    entrySF = tkinter.Entry(sf,bd= 1,width= 12)
    entrySF.place(x=30, y =25)
    lblEnterDEsc= tkinter.Label(sf,text="Enter Desc").place(x=8,y=0)
    btnOk= tkinter.Button(sf,text="Ok",command= lambda:  method3(sf, entrySF.get()))
    btnOk.place(x=40,y=50)
    sf.deiconify()
    sf.mainloop()
    #method1()

def helloCallBackOne():

    global filepath
    top.withdraw()
    top.iconify()
    screenshot= pyautogui.screenshot()
    filepath= entryW.get()
    screenshot.save(filepath+"\screen.png")
    print (filepath)
    global doc
    if os.path.exists(filepath +"\screen.docx")==True:
        print ("exists")
    else:
        doc = None
        doc = docx.Document()
        section = doc.sections[0]
        header = section.header
        t = header.add_table(4, 2, width=Inches(9.0))
        t.style = 'TableGrid'
        tday = date.today()
        strToday = tday.strftime("%d %b %Y")
        t.rows[0].cells[0].text = "Tester"
        t.rows[0].cells[1].text = "Environment"
        t.rows[1].cells[0].text = "Story"
        t.rows[1].cells[1].text = "Timestamp = " + strToday
        t.rows[2].cells[0].text = "User"
        t.rows[2].cells[1].text = "Role"
        t.rows[3].cells[0].text = "Location"
        t.rows[3].cells[1].text = "Browser"

    p = doc.add_paragraph()
    r= p.add_run()
    r.add_break()
    r.add_picture(filepath+"\screen.png", width= Inches(6), height= Inches(4))
    global imageCount
    imageCount = imageCount + 1
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    r.add_text("\t\t\t\t\t Figure "+ str(imageCount))
    font.size = Pt(12)
    doc.save(filepath+"\screen.docx")
    top.deiconify()


def method4():
    t1 = threading.Thread(target=LastResort2.main)
    t1.start()

def method8():
    t2= threading.Thread(target=LastResort3.main)
    t2.start()


def on_press(key):
    try:
        print('alphanumeric key {0} pressed'.format(key.char))
    except AttributeError:
        print('special key {0} pressed'.format(key))

def on_release(key):
    if '{0}'.format(key) == "'`'":
        helloCallBackOne()
    if '{0}'.format(key) == "'-'":
        print("ogothere")
        helloCallBack()
    print('{0} released'.format(key))
    if key == keyboard.Key.esc:
        # Stop listener
        return False

# ...or, in a non-blocking fashion:
listener = keyboard.Listener(
    on_press=on_press,
    on_release=on_release)
listener.start()

#Listener

MyFont = font.Font(size=8)
l = tkinter.Label(top, text= "Path")
l.place(x=0,y=30)
entryW= tkinter.Entry(top,bd=1,width=12)
entryW.place(x=30,y=30)
l['font']= MyFont

B1= tkinter.Button(top,text= "CaptureNText",command= helloCallBack)
B1['font']= MyFont
B1.place(x=54,y=53)
B2= tkinter.Button(top,text= "Capture",command= helloCallBackOne)
B2['font']= MyFont
B2.place(x=1,y=53)
HG= tkinter.Button(top,text= "Highlight",command= method4)
HG.place(x=50,y=2)
HG['font']= MyFont
Help= tkinter.Button(top,text= "?")
Help.place(x=114,y=2)
Help['font']= MyFont
secure= tkinter.Button(top,text= "Secure",command= method8)
secure.place(x=2,y=2)
secure['font']= MyFont

def on_closing():
    if tkinter.messagebox.askokcancel("Quit", "Do you want to quit?"):
        top.destroy()
        os.system("TASKKILL /F /IM CopyApp*")


top.protocol("WM_DELETE_WINDOW", on_closing)
top.mainloop()













