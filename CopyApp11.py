import os
import threading
import tkinter
import tkinter.messagebox
import pyautogui
import docx
from docx.shared import Inches, Pt
import LastResort2
import LastResort3
import LastResort4
from datetime import date
import tkinter.font as font
from pynput import keyboard


global imageCount
imageCount = 0
global textCount
textCount = 0

top = tkinter.Tk()
top.title = "CopyMe"
top.geometry('130x140')
top.resizable(True, True)
top.attributes("-topmost", True)
top.focus()
top['background']='#D6DBDF'
#top['background']='#76D7C4'
#top['background']='#FAD7A0'



global doc
doc = docx.Document()
section = doc.sections[0]
header = section.header
t = header.add_table(4, 2, width=Inches(9.0))
t.style = 'TableGrid'
tday = date.today()
strToday = tday.strftime("%d %b %Y")
t.rows[0].cells[0].text = "Tester"
t.rows[0].cells[1].text = "Environment"
t.rows[1].cells[0].text = "Story"
t.rows[1].cells[1].text = "Timestamp = " + strToday
t.rows[2].cells[0].text = "User"
t.rows[2].cells[1].text = "Role"
t.rows[3].cells[0].text = "Location"
t.rows[3].cells[1].text = "Browser"


def method3(sf,commenttext):
    global doc,filepath,textCount,imageCount
    if os.path.exists(filepath +"\screen.docx")==True:
        print ("exists")
    else:
        textCount = 0
        imageCount=0
        doc = None
        doc = docx.Document()
        section = doc.sections[0]
        header = section.header
        t = header.add_table(4, 2, width=Inches(9.0))
        t.style = 'TableGrid'
        tday = date.today()
        strToday = tday.strftime("%d %b %Y")
        t.rows[0].cells[0].text = "Tester: " + testername
        t.rows[0].cells[1].text = "Environment: "
        t.rows[1].cells[0].text = "Story: " + storyid
        t.rows[1].cells[1].text = "Timestamp: " + strToday
        t.rows[2].cells[0].text = "User: " + userid
        t.rows[2].cells[1].text = "Role: "
        t.rows[3].cells[0].text = "Location: "
        t.rows[3].cells[1].text = "Browser: " + browser

    p = doc.add_paragraph()
    r= p.add_run()
    r.add_break()
    textCount = textCount + 1
    textCountStr = str(textCount) + ".  "
    commenttext = textCountStr + commenttext
    r.add_text(commenttext)
    r.add_picture(filepath+"\screen.png", width= Inches(6), height= Inches(4))
    imageCount = imageCount + 1
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    r.add_text("\t\t\t\t\t Figure "+ str(imageCount))
    doc.save(filepath+"\screen.docx")
    sf.withdraw()
    sf.destroy()
    top.deiconify()

def helloCallBack():
    try:
        print (filepath)
    except NameError:
        tkinter.messagebox.showerror("Error", "Filepath not defined, Click ? and define path")
    else:
    #global filepath
        top.withdraw()
        top.iconify()
        screenshot= pyautogui.screenshot()
        #filepath= entryW.get()
        screenshot.save(filepath+"\screen.png")
        sf = tkinter.Tk()
        sf.geometry('100x80')
        sf.resizable(True,True)
        sf.attributes("-topmost",True)
        entrySF = tkinter.Entry(sf,bd= 1,width= 12)
        entrySF.place(x=30, y =25)
        lblEnterDEsc= tkinter.Label(sf,text="Enter Desc").place(x=8,y=0)
        btnOk= tkinter.Button(sf,text="Ok",command= lambda:  method3(sf, entrySF.get()))
        btnOk.place(x=40,y=50)
        sf.deiconify()
        sf.mainloop()

def method10(sf1,filep,tname,uid,stryid,brwsr):
    global filepath,doc,testername,userid,storyid,browser
    filepath = filep
    print (filepath)
    testername = tname
    userid = uid
    storyid= stryid
    browser = brwsr
    sf1.withdraw()
    sf1.destroy()
    doc= None
    doc = docx.Document()
    top.deiconify

def helloCallBackOne():
    try:
        print (filepath)
    except NameError:
        tkinter.messagebox.showerror("Error","Filepath not defined, Click ? and define path")
    else:
        top.withdraw()
        top.iconify()
        screenshot= pyautogui.screenshot()
        #filepath= entryW.get()
        screenshot.save(filepath+"\screen.png")
        print (filepath)
        global doc,imageCount,textCount
        if os.path.exists(filepath +"\screen.docx")==True:
            print ("exists")
        else:
            textCount=0
            imageCount=0
            doc = None
            doc = docx.Document()
            section = doc.sections[0]
            header = section.header
            t = header.add_table(4, 2, width=Inches(9.0))
            t.style = 'TableGrid'
            tday = date.today()
            strToday = tday.strftime("%d %b %Y")
            t.rows[0].cells[0].text = "Tester: "+ testername
            t.rows[0].cells[1].text = "Environment: "
            t.rows[1].cells[0].text = "Story: "+ storyid
            t.rows[1].cells[1].text = "Timestamp: " + strToday
            t.rows[2].cells[0].text = "User: " + userid
            t.rows[2].cells[1].text = "Role: "
            t.rows[3].cells[0].text = "Location: "
            t.rows[3].cells[1].text = "Browser: "+ browser

        p = doc.add_paragraph()
        r= p.add_run()
        r.add_break()
        r.add_picture(filepath+"\screen.png", width= Inches(6), height= Inches(4))
        imageCount = imageCount + 1
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        r.add_text("\t\t\t\t\t Figure "+ str(imageCount))
        font.size = Pt(12)
        doc.save(filepath+"\screen.docx")
        top.deiconify()

def method4():
    t1 = threading.Thread(target=LastResort2.main)
    t1.start()

def method8():
    t2= threading.Thread(target=LastResort3.main)
    t2.start()

def methodSnip():
    o1 =LastResort4.GraphicsRectItem()
    o1.mainlike()

def method9():
    sf1 = tkinter.Tk()
    sf1.geometry('160x150')
    sf1.resizable(True,True)
    sf1.attributes("-topmost",True)
    sf1['background']='#85929E'
    entrySF1 = tkinter.Entry(sf1,bd= 1,width= 12)
    entrySF1.place(x=75, y =0)
    lblEnterDEsc= tkinter.Label(sf1,text="Path",bg='#85929E').place(x=2,y=0)
    entrySF2 = tkinter.Entry(sf1,bd= 1,width= 12)
    entrySF2.place(x=75, y =20)
    lblEnterDEsc2= tkinter.Label(sf1,text="Tester Name",bg='#85929E').place(x=2,y=20)
    entrySF3 = tkinter.Entry(sf1,bd= 1,width= 12)
    entrySF3.place(x=75, y =40)
    lblEnterDEsc2= tkinter.Label(sf1,text="UserId",bg='#85929E').place(x=2,y=40)
    entrySF4 = tkinter.Entry(sf1,bd= 1,width= 12)
    entrySF4.place(x=75, y =60)
    lblEnterDEsc2= tkinter.Label(sf1,text="Story",bg='#85929E').place(x=2,y=60)
    entrySF5 = tkinter.Entry(sf1,bd= 1,width= 12)
    entrySF5.place(x=75, y =80)
    lblEnterDEsc2= tkinter.Label(sf1,text="Browser",bg='#85929E').place(x=2,y=80)
    btnOk1= tkinter.Button(sf1,text="Ok",command=lambda: method10(sf1,entrySF1.get(),entrySF2.get(),entrySF3.get(),entrySF4.get(),entrySF5.get()))
    btnOk1.place(x=60,y=110)
    sf1.deiconify()
    sf1.mainloop()


def on_press(key):
    try:
        print('alphanumeric key {0} pressed'.format(key.char))
    except AttributeError:
        print('special key {0} pressed'.format(key))

def on_release(key):
    if '{0}'.format(key) == "'`'":
        helloCallBackOne()
    if '{0}'.format(key) == "'~'":
        print("ogothere")
        helloCallBack()
    print('{0} released'.format(key))
    if key == keyboard.Key.f7:
        # Stop listener
        return False

# ...or, in a non-blocking fashion:
listener = keyboard.Listener(
    on_press=on_press,
    on_release=on_release)
listener.start()

#Listener

MyFont = font.Font(size=8)
#l = tkinter.Label(top, text= "Path")
#l.place(x=0,y=30)
#entryW= tkinter.Entry(top,bd=1,width=12)
#entryW.place(x=30,y=30)
#l['font']= MyFont
B1= tkinter.Button(top,text= "CaptureNText",width = 15, height = 1, bg= '#85929E' , command= helloCallBack)

#B1= tkinter.Button(top,text= "CaptureNText",width = 15, height = 1, bg= '#F8C471' , command= helloCallBack)
#B1= tkinter.Button(top,text= "CaptureNText",width = 15, height = 1, bg= '#73C6B6' , command= helloCallBack)
B1['font']= MyFont
B1.place(x=1,y=5)
B2= tkinter.Button(top,text= "Capture",width = 15, height = 1, bg = '#5D6D7E', command= helloCallBackOne)
#B2= tkinter.Button(top,text= "Capture",width = 15, height = 1, bg = '#F5B041', command= helloCallBackOne)
#B2= tkinter.Button(top,text= "Capture",width = 15, height = 1, bg = '#45B39D', command= helloCallBackOne)
B2['font']= MyFont
B2.place(x=1,y=30)
HG= tkinter.Button(top,text= "Highlight",width = 15, height = 1, bg= '#34495E', command= method4)
HG.place(x=1,y=55)
HG['font']= MyFont
SC= tkinter.Button(top,text= "Snip",width = 15, height = 1, bg= '#34495E', command= methodSnip)
SC.place(x=1,y=105)
SC['font']= MyFont
Help= tkinter.Button(top,text= "?",width = 1, height = 1, bg='#85929E',command = method9)
Help.place(x=114,y=2)
Help['font']= MyFont
secure= tkinter.Button(top,text= "Mask",width = 15,bg='#2E4053', height = 1,command= method8)
#secure= tkinter.Button(top,text= "Secure",width = 15,bg='#D68910', height = 1,command= method8)
#secure= tkinter.Button(top,text= "Secure",width = 15,bg='#138D75', height = 1,command= method8)
secure.place(x=1,y=80)
secure['font']= MyFont

def on_closing():
    if tkinter.messagebox.askokcancel("Quit", "Do you want to quit?"):
        top.destroy()
        os.system("TASKKILL /F /IM CopyApp*")
        os.system("TASKKILL /F /IM python*")


top.protocol("WM_DELETE_WINDOW", on_closing)
top.mainloop()













